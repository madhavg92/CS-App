from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Query
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
from emergentintegrations.llm.chat import LlmChat, UserMessage
import re
import jwt
import io
import csv
import json
import bcrypt
import asyncio
from fastapi.responses import StreamingResponse
import openpyxl

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI(title="Anka CS Hub API", version="2.0.0")
api_router = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)

JWT_SECRET = os.environ.get('JWT_SECRET', 'anka-cs-hub-secret-key-change-in-production')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# =============================================================================
# PYDANTIC MODELS - Production Data Structure
# =============================================================================

# User & Auth Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: str
    password_hash: str
    role: str  # cs_lead, csm, ops
    assigned_clients: List[str] = []
    is_active: bool = True
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    last_login: Optional[str] = None

class UserCreate(BaseModel):
    name: str
    email: EmailStr
    password: str
    role: str = "csm"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    name: str
    email: str
    role: str
    assigned_clients: List[str]
    is_active: bool

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

# Client Account Model
class ClientAccount(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    contract_start: str
    contract_end: str
    contracted_services: List[str] = []  # EV, Prior Auth, Coding, Billing, AR, Payment Posting
    sla_targets: Dict[str, Any] = {}  # {"recovery_rate": 85, "error_rate": 2, "turnaround_days": 3}
    status: str = "active"  # active, at-risk, churned
    assigned_csm: Optional[str] = None
    assigned_cs_lead: Optional[str] = None
    health_score: int = 80  # Computed field
    last_contact_date: Optional[str] = None
    last_innovation_briefing: Optional[str] = None  # ISO date string for innovation update tracking
    notes: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ClientAccountCreate(BaseModel):
    name: str
    contract_start: str
    contract_end: str
    contracted_services: List[str] = []
    sla_targets: Dict[str, Any] = {}
    assigned_csm: Optional[str] = None
    assigned_cs_lead: Optional[str] = None
    last_innovation_briefing: Optional[str] = None

# Client Contact Model
class ClientContact(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    name: str
    title: str
    email: str
    phone: Optional[str] = None
    role_type: str  # decision-maker, influencer, operations, billing
    anka_relationship_owner: Optional[str] = None
    last_email_date: Optional[str] = None
    last_call_date: Optional[str] = None
    last_meeting_date: Optional[str] = None
    notes: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ClientContactCreate(BaseModel):
    client_id: str
    name: str
    title: str
    email: EmailStr
    phone: Optional[str] = None
    role_type: str
    anka_relationship_owner: Optional[str] = None

# Performance Record Model
class PerformanceRecord(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    period_start: str
    period_end: str
    denials_worked: int
    dollars_recovered: float
    recovery_rate: float
    sla_compliance_pct: float
    error_rate: float
    top_denial_codes: Dict[str, int] = {}  # {"CO-4": 150, "PR-1": 89}
    payer_breakdown: Dict[str, float] = {}  # {"Medicare": 45000, "Medicaid": 23000}
    claims_processed: Optional[int] = None  # Team productivity
    avg_turnaround_days: Optional[float] = None  # Team productivity
    team_size: Optional[int] = None  # Team productivity
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class PerformanceRecordCreate(BaseModel):
    client_id: str
    period_start: str
    period_end: str
    denials_worked: int
    dollars_recovered: float
    recovery_rate: float
    sla_compliance_pct: float
    error_rate: float
    top_denial_codes: Dict[str, int] = {}
    payer_breakdown: Dict[str, float] = {}
    claims_processed: Optional[int] = None
    avg_turnaround_days: Optional[float] = None
    team_size: Optional[int] = None

# Alert Model
class Alert(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    alert_type: str  # engagement_gap, frustration, scope_creep, new_stakeholder, policy_update, renewal, performance_decline, followup_overdue
    severity: str  # critical, high, medium, low
    title: str
    description: str
    trigger_data: Dict[str, Any] = {}
    status: str = "active"  # active, acknowledged, resolved, snoozed
    assigned_to: Optional[str] = None
    snoozed_until: Optional[str] = None
    resolution_notes: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    acknowledged_at: Optional[str] = None
    resolved_at: Optional[str] = None

class AlertCreate(BaseModel):
    client_id: str
    client_name: str
    alert_type: str
    severity: str
    title: str
    description: str
    trigger_data: Dict[str, Any] = {}
    assigned_to: Optional[str] = None

class AlertUpdate(BaseModel):
    status: Optional[str] = None
    assigned_to: Optional[str] = None
    snoozed_until: Optional[str] = None
    resolution_notes: Optional[str] = None

# Communication Model
class Communication(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    contact_id: Optional[str] = None
    comm_type: str  # draft, sent, received
    channel: str  # email, call, meeting
    subject: str
    body: str
    ai_generated: bool = False
    template_used: Optional[str] = None
    created_by: Optional[str] = None
    sent_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CommunicationCreate(BaseModel):
    client_id: str
    contact_id: Optional[str] = None
    comm_type: str
    channel: str
    subject: str
    body: str
    ai_generated: bool = False
    template_used: Optional[str] = None

# Follow-Up Item Model
class FollowUpItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    description: str
    days_open: int = 0
    owner: str
    priority_score: float = 5.0  # Computed: 1-10
    priority_level: int = 3  # 1-5
    category: str  # call_required, email_sufficient
    suggested_action: str
    status: str = "open"  # open, in_progress, completed
    draft_email_id: Optional[str] = None
    source: str = "manual"  # manual, crm, alert, meeting
    due_date: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    completed_at: Optional[str] = None

class FollowUpCreate(BaseModel):
    client_id: str
    client_name: str
    description: str
    owner: str
    priority_level: int = 3
    suggested_action: str
    source: str = "manual"
    due_date: Optional[str] = None

# Prompt Template Model
class PromptTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    slug: str  # weekly_summary, client_report_external, followup_list, engagement_outreach, qbr_scheduling, escalation_notification
    description: str
    system_prompt: str
    user_prompt_template: str
    parameters: List[str] = []  # Required parameters
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Alert Threshold Configuration
class AlertThreshold(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: Optional[str] = None  # None = global default
    engagement_gap_decision_maker_days: int = 14
    engagement_gap_influencer_days: int = 30
    renewal_alert_60_days: bool = True
    renewal_alert_30_days: bool = True
    renewal_alert_15_days: bool = True
    performance_decline_threshold_pct: float = 5.0
    followup_overdue_days: int = 7
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# CRM Sync Models (HubSpot-like structure)
class CRMContact(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    firstname: str
    lastname: str
    phone: Optional[str] = None
    company: Optional[str] = None
    jobtitle: Optional[str] = None
    lifecyclestage: str = "lead"
    hs_lead_status: Optional[str] = None
    createdate: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    lastmodifieddate: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CRMDeal(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    dealname: str
    amount: float
    dealstage: str  # appointmentscheduled, qualifiedtobuy, presentationscheduled, closedwon, closedlost
    pipeline: str = "default"
    closedate: Optional[str] = None
    associated_contact_ids: List[str] = []
    createdate: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    lastmodifieddate: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CRMTicket(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    subject: str
    content: str
    hs_pipeline: str = "support"
    hs_pipeline_stage: str = "new"  # new, waiting, closed
    hs_ticket_priority: str = "medium"
    associated_contact_ids: List[str] = []
    createdate: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# PHI Redaction Log
class PHIRedactionLog(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    original_field: str
    redaction_type: str  # SSN, DOB, MRN, PHONE
    context: str
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Audit Log Model
class AuditLog(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    user_name: str
    action_type: str  # ai_generation, report_download, alert_action, communication_send, data_upload
    resource_type: str
    resource_id: Optional[str] = None
    details: Dict[str, Any] = {}
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Scheduled Review Model (for Calendar)
class ScheduledReview(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    review_type: str  # weekly, monthly, qbr
    scheduled_date: str
    scheduled_time: str
    duration_minutes: int = 60
    attendee_ids: List[str] = []
    attendees: List[str] = []  # Names for display
    status: str = "scheduled"  # scheduled, completed, cancelled
    notes: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Team Member Model (for Org Mapping)
class TeamMember(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    role: str  # manager, supervisor, team_lead
    email: str
    phone: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# LOB Model
class LOB(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Client LOB Mapping
class ClientLOBMapping(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    lob_id: str
    lob_name: str
    anka_team_members: List[str] = []  # team member IDs
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Integration Config
class IntegrationConfig(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    integration_name: str
    connection_status: str = "disconnected"  # connected, disconnected, error
    config: Dict[str, Any] = {}
    last_sync: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Policy Update Model (for policy alert system)
class PolicyUpdate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: str
    policy_type: str  # payer_update, regulatory, compliance, billing_rule
    affected_services: List[str] = []  # EV, Prior Auth, Coding, Billing, AR, Payment Posting
    affected_payers: List[str] = []  # Medicare, Medicaid, BCBS, United, Aetna, Cigna
    effective_date: Optional[str] = None
    source_url: Optional[str] = None
    status: str = "active"  # active, archived
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Email Cadence Model
class EmailCadence(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    cadence_type: str  # monthly_checkin, qbr_scheduling, renewal_reminder, open_items_followup
    frequency_days: int = 30
    auto_send: bool = False
    status: str = "active"  # active, paused
    last_triggered: Optional[str] = None
    next_trigger: Optional[str] = None
    template_slug: str = "engagement_outreach"
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def hash_password(password: str) -> str:
    """Hash password using bcrypt"""
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, password_hash: str) -> bool:
    """Verify password against bcrypt hash"""
    try:
        return bcrypt.checkpw(password.encode('utf-8'), password_hash.encode('utf-8'))
    except Exception:
        # Fallback for legacy SHA-256 hashes during migration
        import hashlib
        return hashlib.sha256(password.encode()).hexdigest() == password_hash

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[dict]:
    if not credentials:
        return None
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def require_auth(user: dict = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required")
    return user

def require_role(allowed_roles: List[str]):
    def role_checker(user: dict = Depends(require_auth)):
        if user["role"] not in allowed_roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

# PHI Scrubbing Middleware
PHI_PATTERNS = {
    "SSN": r'\b\d{3}-\d{2}-\d{4}\b',
    "DOB": r'\b(0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])[-/](19|20)\d{2}\b',
    "MRN": r'\b(MRN|mrn|Medical Record|medical record)[:\s#]*\d{5,12}\b',
    "PHONE_CLINICAL": r'\b(patient|pt|member)[\s\w]*phone[:\s]*\d{3}[-.]?\d{3}[-.]?\d{4}\b',
}

def scrub_phi(text: str, context: str = "unknown") -> tuple[str, List[dict]]:
    """Scrub PHI from text and return (scrubbed_text, redaction_logs)"""
    redactions = []
    scrubbed = text
    
    for phi_type, pattern in PHI_PATTERNS.items():
        matches = re.findall(pattern, scrubbed, re.IGNORECASE)
        if matches:
            scrubbed = re.sub(pattern, f'[REDACTED-{phi_type}]', scrubbed, flags=re.IGNORECASE)
            for _ in matches:
                redactions.append({
                    "original_field": context,
                    "redaction_type": phi_type,
                    "context": context
                })
    
    return scrubbed, redactions

async def log_phi_redactions(redactions: List[dict]):
    """Log PHI redactions to database (without original values)"""
    if redactions:
        for r in redactions:
            log = PHIRedactionLog(**r)
            await db.phi_redaction_logs.insert_one(log.model_dump())

async def log_audit(user_id: str, user_name: str, action_type: str, resource_type: str, resource_id: str = None, details: dict = None):
    """Log an audit entry"""
    audit = AuditLog(
        user_id=user_id,
        user_name=user_name,
        action_type=action_type,
        resource_type=resource_type,
        resource_id=resource_id,
        details=details or {}
    )
    await db.audit_logs.insert_one(audit.model_dump())

# AI Helper
async def generate_with_ai(prompt: str, system_message: str = "You are a helpful AI assistant for healthcare customer success operations. Be professional, accurate, and reference specific data when provided.") -> str:
    try:
        api_key = os.environ.get('EMERGENT_LLM_KEY')
        chat = LlmChat(
            api_key=api_key,
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        return response
    except Exception as e:
        logging.error(f"AI generation error: {e}")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")

# Health Score Calculator
async def calculate_health_score(client_id: str) -> int:
    """Calculate health score based on engagement, SLA compliance, and open issues"""
    score = 100
    
    # Get client
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        return 80
    
    # Factor 1: Engagement recency (30% weight)
    contacts = await db.client_contacts.find({"client_id": client_id}, {"_id": 0}).to_list(100)
    if contacts:
        most_recent = None
        for contact in contacts:
            for date_field in ['last_email_date', 'last_call_date', 'last_meeting_date']:
                if contact.get(date_field):
                    if not most_recent or contact[date_field] > most_recent:
                        most_recent = contact[date_field]
        
        if most_recent:
            days_since = (datetime.now(timezone.utc) - datetime.fromisoformat(most_recent.replace('Z', '+00:00'))).days
            if days_since > 30:
                score -= 20
            elif days_since > 14:
                score -= 10
    else:
        score -= 15
    
    # Factor 2: SLA Compliance (40% weight)
    recent_perf = await db.performance_records.find(
        {"client_id": client_id}
    ).sort("period_end", -1).limit(1).to_list(1)
    
    if recent_perf:
        perf = recent_perf[0]
        sla_target = client.get('sla_targets', {}).get('recovery_rate', 85)
        if perf.get('recovery_rate', 0) < sla_target - 10:
            score -= 25
        elif perf.get('recovery_rate', 0) < sla_target:
            score -= 15
        
        if perf.get('sla_compliance_pct', 100) < 90:
            score -= 15
    
    # Factor 3: Open Issues (30% weight)
    open_alerts = await db.alerts.count_documents({"client_id": client_id, "status": {"$in": ["active", "acknowledged"]}})
    if open_alerts >= 3:
        score -= 20
    elif open_alerts >= 1:
        score -= 10
    
    open_followups = await db.followups.count_documents({"client_id": client_id, "status": "open"})
    if open_followups >= 5:
        score -= 10
    elif open_followups >= 2:
        score -= 5
    
    return max(0, min(100, score))

# Priority Score Calculator for Follow-ups
async def calculate_priority_score(followup: dict) -> float:
    """Calculate priority score based on days_open, client_health, severity, renewal proximity"""
    score = 0.0
    
    # Days open (weight: 0.3)
    days_open = followup.get('days_open', 0)
    days_score = min(10, days_open / 3)  # Max out at 30 days
    score += days_score * 0.3
    
    # Client health score (weight: 0.3) - lower health = higher priority
    client = await db.clients.find_one({"id": followup.get('client_id')}, {"_id": 0})
    if client:
        health = client.get('health_score', 80)
        health_score = (100 - health) / 10  # Convert to 0-10 scale
        score += health_score * 0.3
        
        # Renewal proximity (weight: 0.2)
        if client.get('contract_end'):
            try:
                contract_end_str = client['contract_end']
                # Handle various date formats
                if 'T' in contract_end_str or '+' in contract_end_str or 'Z' in contract_end_str:
                    contract_end = datetime.fromisoformat(contract_end_str.replace('Z', '+00:00'))
                else:
                    # Simple date string like "2024-01-01"
                    contract_end = datetime.fromisoformat(contract_end_str + "T00:00:00+00:00")
                days_to_renewal = (contract_end - datetime.now(timezone.utc)).days
                if days_to_renewal <= 30:
                    score += 10 * 0.2
                elif days_to_renewal <= 60:
                    score += 7 * 0.2
                elif days_to_renewal <= 90:
                    score += 4 * 0.2
            except Exception:
                pass  # Skip renewal calculation if date parsing fails
    
    # Priority level (weight: 0.2)
    priority_level = followup.get('priority_level', 3)
    score += (priority_level * 2) * 0.2
    
    return round(score, 2)

# =============================================================================
# AUTH ROUTES
# =============================================================================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(input: UserCreate):
    existing = await db.users.find_one({"email": input.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user = User(
        name=input.name,
        email=input.email,
        password_hash=hash_password(input.password),
        role=input.role
    )
    await db.users.insert_one(user.model_dump())
    
    token = create_token(user.id, user.email, user.role)
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user.id,
            name=user.name,
            email=user.email,
            role=user.role,
            assigned_clients=user.assigned_clients,
            is_active=user.is_active
        )
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(input: UserLogin):
    user = await db.users.find_one({"email": input.email}, {"_id": 0})
    if not user or not verify_password(input.password, user['password_hash']):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not user.get('is_active', True):
        raise HTTPException(status_code=403, detail="Account disabled")
    
    # Update last login
    await db.users.update_one(
        {"id": user['id']},
        {"$set": {"last_login": datetime.now(timezone.utc).isoformat()}}
    )
    
    token = create_token(user['id'], user['email'], user['role'])
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user['id'],
            name=user['name'],
            email=user['email'],
            role=user['role'],
            assigned_clients=user.get('assigned_clients', []),
            is_active=user.get('is_active', True)
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user: dict = Depends(require_auth)):
    return UserResponse(
        id=user['id'],
        name=user['name'],
        email=user['email'],
        role=user['role'],
        assigned_clients=user.get('assigned_clients', []),
        is_active=user.get('is_active', True)
    )

# =============================================================================
# USER MANAGEMENT (cs_lead only)
# =============================================================================

@api_router.get("/users", response_model=List[UserResponse])
async def get_users(user: dict = Depends(require_role(["cs_lead"]))):
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return [UserResponse(**u) for u in users]

@api_router.patch("/users/{user_id}")
async def update_user(user_id: str, role: Optional[str] = None, assigned_clients: Optional[List[str]] = None, is_active: Optional[bool] = None, admin: dict = Depends(require_role(["cs_lead"]))):
    update_data = {}
    if role:
        update_data["role"] = role
    if assigned_clients is not None:
        update_data["assigned_clients"] = assigned_clients
    if is_active is not None:
        update_data["is_active"] = is_active
    
    if update_data:
        result = await db.users.update_one({"id": user_id}, {"$set": update_data})
        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User updated"}

# =============================================================================
# CLIENT ROUTES
# =============================================================================

@api_router.get("/clients", response_model=List[ClientAccount])
async def get_clients(user: dict = Depends(require_auth)):
    query = {}
    # CSMs can only see their assigned clients
    if user['role'] == 'csm':
        query["id"] = {"$in": user.get('assigned_clients', [])}
    
    clients = await db.clients.find(query, {"_id": 0}).to_list(1000)
    
    # Update health scores
    for client in clients:
        client['health_score'] = await calculate_health_score(client['id'])
    
    return clients

@api_router.get("/clients/{client_id}", response_model=ClientAccount)
async def get_client(client_id: str, user: dict = Depends(require_auth)):
    # Check access
    if user['role'] == 'csm' and client_id not in user.get('assigned_clients', []):
        raise HTTPException(status_code=403, detail="Access denied to this client")
    
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    client['health_score'] = await calculate_health_score(client_id)
    return client

@api_router.post("/clients", response_model=ClientAccount)
async def create_client(input: ClientAccountCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    client = ClientAccount(
        name=input.name,
        contract_start=input.contract_start,
        contract_end=input.contract_end,
        contracted_services=input.contracted_services,
        sla_targets=input.sla_targets,
        assigned_csm=input.assigned_csm,
        assigned_cs_lead=input.assigned_cs_lead
    )
    await db.clients.insert_one(client.model_dump())
    return client

@api_router.put("/clients/{client_id}", response_model=ClientAccount)
async def update_client(client_id: str, input: ClientAccountCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    update_data = input.model_dump()
    update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
    
    result = await db.clients.update_one({"id": client_id}, {"$set": update_data})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return await get_client(client_id, user)

# =============================================================================
# CLIENT CONTACTS ROUTES
# =============================================================================

@api_router.get("/contacts", response_model=List[ClientContact])
async def get_contacts(client_id: Optional[str] = None, user: dict = Depends(require_auth)):
    query = {}
    if client_id:
        query["client_id"] = client_id
    
    contacts = await db.client_contacts.find(query, {"_id": 0}).to_list(1000)
    return contacts

@api_router.post("/contacts", response_model=ClientContact)
async def create_contact(input: ClientContactCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    contact = ClientContact(**input.model_dump())
    await db.client_contacts.insert_one(contact.model_dump())
    return contact

@api_router.put("/contacts/{contact_id}", response_model=ClientContact)
async def update_contact(contact_id: str, input: ClientContactCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    result = await db.client_contacts.update_one({"id": contact_id}, {"$set": input.model_dump()})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Contact not found")
    
    contact = await db.client_contacts.find_one({"id": contact_id}, {"_id": 0})
    return contact

# =============================================================================
# PERFORMANCE RECORDS ROUTES
# =============================================================================

@api_router.get("/performance", response_model=List[PerformanceRecord])
async def get_performance_records(client_id: Optional[str] = None, user: dict = Depends(require_auth)):
    query = {}
    if client_id:
        query["client_id"] = client_id
    
    records = await db.performance_records.find(query, {"_id": 0}).sort("period_end", -1).to_list(1000)
    return records

@api_router.post("/performance", response_model=PerformanceRecord)
async def create_performance_record(input: PerformanceRecordCreate, user: dict = Depends(require_role(["cs_lead", "csm", "ops"]))):
    record = PerformanceRecord(**input.model_dump())
    await db.performance_records.insert_one(record.model_dump())
    
    # Check for performance decline alert
    await check_performance_decline_alert(input.client_id, record)
    
    return record

async def check_performance_decline_alert(client_id: str, new_record: PerformanceRecord):
    """Check if performance declined and create alert if needed"""
    # Get previous record
    prev_records = await db.performance_records.find(
        {"client_id": client_id, "id": {"$ne": new_record.id}}
    ).sort("period_end", -1).limit(1).to_list(1)
    
    if not prev_records:
        return
    
    prev = prev_records[0]
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    client_name = client.get('name', 'Unknown') if client else 'Unknown'
    
    # Get threshold
    threshold = await db.alert_thresholds.find_one({"client_id": client_id}, {"_id": 0})
    if not threshold:
        threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
    decline_threshold = threshold.get('performance_decline_threshold_pct', 5.0) if threshold else 5.0
    
    # Check recovery rate decline
    if prev.get('recovery_rate', 0) - new_record.recovery_rate > decline_threshold:
        alert = Alert(
            client_id=client_id,
            client_name=client_name,
            alert_type="performance_decline",
            severity="high",
            title=f"Recovery Rate Declined by {prev.get('recovery_rate', 0) - new_record.recovery_rate:.1f}%",
            description=f"Recovery rate dropped from {prev.get('recovery_rate', 0):.1f}% to {new_record.recovery_rate:.1f}%",
            trigger_data={
                "previous_rate": prev.get('recovery_rate', 0),
                "current_rate": new_record.recovery_rate,
                "decline_pct": prev.get('recovery_rate', 0) - new_record.recovery_rate
            }
        )
        await db.alerts.insert_one(alert.model_dump())
    
    # Check SLA compliance
    sla_target = client.get('sla_targets', {}).get('sla_compliance', 95) if client else 95
    if new_record.sla_compliance_pct < sla_target:
        alert = Alert(
            client_id=client_id,
            client_name=client_name,
            alert_type="performance_decline",
            severity="medium",
            title="SLA Compliance Below Target",
            description=f"SLA compliance at {new_record.sla_compliance_pct:.1f}% vs target of {sla_target}%",
            trigger_data={
                "current_compliance": new_record.sla_compliance_pct,
                "target": sla_target
            }
        )
        await db.alerts.insert_one(alert.model_dump())

# OPS DATA UPLOAD
@api_router.post("/performance/upload")
async def upload_ops_data(file: UploadFile = File(...), user: dict = Depends(require_role(["cs_lead", "ops"]))):
    """Upload CSV/Excel file with ops data"""
    if not file.filename.endswith(('.csv', '.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only CSV and Excel files are supported")
    
    content = await file.read()
    rows = []
    
    try:
        # Parse CSV or Excel
        if file.filename.endswith('.csv'):
            decoded = content.decode('utf-8')
            reader = csv.DictReader(io.StringIO(decoded))
            rows = list(reader)
        else:
            # Parse Excel file
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            sheet = workbook.active
            headers = [cell.value for cell in sheet[1]]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if any(row):  # Skip empty rows
                    row_dict = {headers[i]: row[i] for i in range(len(headers)) if headers[i]}
                    rows.append(row_dict)
        
        required_columns = ['client_name', 'period_start', 'period_end', 'denials_worked', 'dollars_recovered', 'recovery_rate', 'sla_compliance', 'error_rate']
        
        # Validate columns
        if rows:
            missing = [col for col in required_columns if col not in rows[0].keys()]
            if missing:
                raise HTTPException(status_code=400, detail=f"Missing columns: {', '.join(missing)}")
        
        records_created = 0
        errors = []
        
        for i, row in enumerate(rows):
            try:
                # Find client by name
                client = await db.clients.find_one({"name": row['client_name']}, {"_id": 0})
                if not client:
                    errors.append(f"Row {i+1}: Client '{row['client_name']}' not found")
                    continue
                
                # Parse denial codes and payer breakdown if present
                top_denial_codes = {}
                payer_breakdown = {}
                
                if row.get('top_denial_codes'):
                    try:
                        if isinstance(row['top_denial_codes'], str):
                            top_denial_codes = json.loads(row['top_denial_codes'])
                        elif isinstance(row['top_denial_codes'], dict):
                            top_denial_codes = row['top_denial_codes']
                    except:
                        pass
                
                if row.get('payer_breakdown'):
                    try:
                        if isinstance(row['payer_breakdown'], str):
                            payer_breakdown = json.loads(row['payer_breakdown'])
                        elif isinstance(row['payer_breakdown'], dict):
                            payer_breakdown = row['payer_breakdown']
                    except:
                        pass
                
                record = PerformanceRecord(
                    client_id=client['id'],
                    period_start=str(row['period_start']),
                    period_end=str(row['period_end']),
                    denials_worked=int(float(row['denials_worked'])),
                    dollars_recovered=float(row['dollars_recovered']),
                    recovery_rate=float(row['recovery_rate']),
                    sla_compliance_pct=float(row['sla_compliance']),
                    error_rate=float(row['error_rate']),
                    top_denial_codes=top_denial_codes,
                    payer_breakdown=payer_breakdown
                )
                
                await db.performance_records.insert_one(record.model_dump())
                await check_performance_decline_alert(client['id'], record)
                records_created += 1
                
            except Exception as e:
                errors.append(f"Row {i+1}: {str(e)}")
        
        # Log audit
        await log_audit(user['id'], user['name'], 'data_upload', 'performance_records', None, {
            'file_name': file.filename,
            'records_created': records_created,
            'errors_count': len(errors)
        })
        
        return {
            "message": f"Uploaded {records_created} records",
            "records_created": records_created,
            "errors": errors
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

# =============================================================================
# ALERTS ROUTES
# =============================================================================

@api_router.get("/alerts", response_model=List[Alert])
async def get_alerts(
    status: Optional[str] = None,
    alert_type: Optional[str] = None,
    severity: Optional[str] = None,
    assigned_to: Optional[str] = None,
    client_id: Optional[str] = None,
    user: dict = Depends(require_auth)
):
    query = {}
    if status:
        query["status"] = status
    if alert_type:
        query["alert_type"] = alert_type
    if severity:
        query["severity"] = severity
    if assigned_to:
        query["assigned_to"] = assigned_to
    if client_id:
        query["client_id"] = client_id
    
    # CSMs only see alerts for their clients
    if user['role'] == 'csm':
        query["client_id"] = {"$in": user.get('assigned_clients', [])}
    
    alerts = await db.alerts.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return alerts

@api_router.get("/alerts/today", response_model=List[Alert])
async def get_todays_alerts(user: dict = Depends(require_auth)):
    """Get alerts created today, sorted by severity"""
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    
    query = {"created_at": {"$gte": today_start}, "status": {"$in": ["active", "acknowledged"]}}
    if user['role'] == 'csm':
        query["client_id"] = {"$in": user.get('assigned_clients', [])}
    
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    alerts = await db.alerts.find(query, {"_id": 0}).to_list(1000)
    alerts.sort(key=lambda x: severity_order.get(x.get('severity', 'low'), 3))
    
    return alerts

@api_router.post("/alerts", response_model=Alert)
async def create_alert(input: AlertCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    alert = Alert(**input.model_dump())
    await db.alerts.insert_one(alert.model_dump())
    return alert

@api_router.patch("/alerts/{alert_id}")
async def update_alert(alert_id: str, input: AlertUpdate, user: dict = Depends(require_auth)):
    update_data = {k: v for k, v in input.model_dump().items() if v is not None}
    
    if input.status == "acknowledged":
        update_data["acknowledged_at"] = datetime.now(timezone.utc).isoformat()
    elif input.status == "resolved":
        update_data["resolved_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.alerts.update_one({"id": alert_id}, {"$set": update_data})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Alert not found")
    
    # Audit log for alert action
    await log_audit(user['id'], user['name'], 'alert_action', 'alert', alert_id, {'action': update_data.get('status'), 'notes': update_data.get('resolution_notes')})
    
    return {"message": "Alert updated"}

@api_router.post("/alerts/{alert_id}/snooze")
async def snooze_alert(alert_id: str, days: int = 1, user: dict = Depends(require_auth)):
    snooze_until = (datetime.now(timezone.utc) + timedelta(days=days)).isoformat()
    result = await db.alerts.update_one(
        {"id": alert_id},
        {"$set": {"status": "snoozed", "snoozed_until": snooze_until}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Alert not found")
    
    # Audit log for snooze
    await log_audit(user['id'], user['name'], 'alert_action', 'alert', alert_id, {'action': 'snoozed', 'days': days})
    
    return {"message": f"Alert snoozed for {days} day(s)", "snoozed_until": snooze_until}

@api_router.post("/alerts/{alert_id}/escalate")
async def escalate_alert(alert_id: str, escalate_to: str, notes: str, user: dict = Depends(require_auth)):
    alert = await db.alerts.find_one({"id": alert_id}, {"_id": 0})
    if not alert:
        raise HTTPException(status_code=404, detail="Alert not found")
    
    await db.alerts.update_one(
        {"id": alert_id},
        {"$set": {
            "severity": "critical",
            "assigned_to": escalate_to,
            "trigger_data": {**alert.get('trigger_data', {}), "escalation_notes": notes, "escalated_by": user['name']}
        }}
    )
    
    # Audit log for escalation
    await log_audit(user['id'], user['name'], 'alert_action', 'alert', alert_id, {'action': 'escalate', 'escalated_to': escalate_to})
    
    return {"message": "Alert escalated"}

# =============================================================================
# COMMUNICATIONS ROUTES
# =============================================================================

@api_router.get("/communications", response_model=List[Communication])
async def get_communications(
    client_id: Optional[str] = None,
    comm_type: Optional[str] = None,
    user: dict = Depends(require_auth)
):
    # Ops cannot access communications
    if user['role'] == 'ops':
        raise HTTPException(status_code=403, detail="Communications access not allowed for ops role")
    
    query = {}
    if client_id:
        query["client_id"] = client_id
    if comm_type:
        query["comm_type"] = comm_type
    
    comms = await db.communications.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return comms

@api_router.get("/communications/drafts", response_model=List[Communication])
async def get_draft_communications(user: dict = Depends(require_auth)):
    if user['role'] == 'ops':
        raise HTTPException(status_code=403, detail="Communications access not allowed")
    
    drafts = await db.communications.find({"comm_type": "draft"}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return drafts

@api_router.post("/communications", response_model=Communication)
async def create_communication(input: CommunicationCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    # Scrub PHI from body
    scrubbed_body, redactions = scrub_phi(input.body, f"communication_{input.client_id}")
    await log_phi_redactions(redactions)
    
    comm = Communication(
        client_id=input.client_id,
        contact_id=input.contact_id,
        comm_type=input.comm_type,
        channel=input.channel,
        subject=input.subject,
        body=scrubbed_body,
        ai_generated=input.ai_generated,
        template_used=input.template_used,
        created_by=user['id']
    )
    await db.communications.insert_one(comm.model_dump())
    
    # Update contact's last communication date
    if input.contact_id and input.comm_type == "sent":
        date_field = f"last_{input.channel}_date"
        await db.client_contacts.update_one(
            {"id": input.contact_id},
            {"$set": {date_field: datetime.now(timezone.utc).isoformat()}}
        )
    
    # Audit log
    await log_audit(user['id'], user['name'], 'communication_send', 'communication', comm.id, {'channel': comm.channel, 'ai_generated': comm.ai_generated})
    
    return comm

@api_router.patch("/communications/{comm_id}/send")
async def send_communication(comm_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    result = await db.communications.update_one(
        {"id": comm_id},
        {"$set": {"comm_type": "sent", "sent_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Communication not found")
    
    # Audit log
    await log_audit(user['id'], user['name'], 'communication_send', 'communication', comm_id, {'action': 'sent'})
    
    return {"message": "Communication sent"}

# =============================================================================
# FOLLOW-UPS ROUTES
# =============================================================================

@api_router.get("/followups", response_model=List[FollowUpItem])
async def get_followups(
    status: Optional[str] = None,
    client_id: Optional[str] = None,
    owner: Optional[str] = None,
    user: dict = Depends(require_auth)
):
    query = {}
    if status:
        query["status"] = status
    if client_id:
        query["client_id"] = client_id
    if owner:
        query["owner"] = owner
    
    followups = await db.followups.find(query, {"_id": 0}).to_list(1000)
    
    # Update days_open and priority_score
    now = datetime.now(timezone.utc)
    for f in followups:
        created = datetime.fromisoformat(f['created_at'].replace('Z', '+00:00'))
        f['days_open'] = (now - created).days
        f['priority_score'] = await calculate_priority_score(f)
    
    # Sort by priority score descending
    followups.sort(key=lambda x: x.get('priority_score', 0), reverse=True)
    
    return followups

@api_router.get("/followups/call-list", response_model=List[FollowUpItem])
async def get_call_list(user: dict = Depends(require_auth)):
    """Get today's prioritized call list (top 10 call_required items)"""
    query = {"status": "open", "category": "call_required"}
    if user['role'] == 'csm':
        query["owner"] = user['id']
    
    followups = await db.followups.find(query, {"_id": 0}).to_list(100)
    
    # Update and sort
    for f in followups:
        f['priority_score'] = await calculate_priority_score(f)
    
    followups.sort(key=lambda x: x.get('priority_score', 0), reverse=True)
    
    return followups[:10]

@api_router.post("/followups", response_model=FollowUpItem)
async def create_followup(input: FollowUpCreate, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    # Determine category based on priority
    category = "call_required" if input.priority_level >= 4 else "email_sufficient"
    
    followup = FollowUpItem(
        client_id=input.client_id,
        client_name=input.client_name,
        description=input.description,
        owner=input.owner,
        priority_level=input.priority_level,
        category=category,
        suggested_action=input.suggested_action,
        source=input.source,
        due_date=input.due_date
    )
    
    followup.priority_score = await calculate_priority_score(followup.model_dump())
    
    await db.followups.insert_one(followup.model_dump())
    return followup

@api_router.patch("/followups/{followup_id}")
async def update_followup(followup_id: str, status: Optional[str] = None, owner: Optional[str] = None, user: dict = Depends(require_auth)):
    update_data = {}
    if status:
        update_data["status"] = status
        if status == "completed":
            update_data["completed_at"] = datetime.now(timezone.utc).isoformat()
    if owner:
        update_data["owner"] = owner
    
    if update_data:
        result = await db.followups.update_one({"id": followup_id}, {"$set": update_data})
        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="Follow-up not found")
    
    return {"message": "Follow-up updated"}

@api_router.post("/followups/{followup_id}/generate-draft")
async def generate_followup_draft(followup_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Generate AI email draft for a follow-up item"""
    followup = await db.followups.find_one({"id": followup_id}, {"_id": 0})
    if not followup:
        raise HTTPException(status_code=404, detail="Follow-up not found")
    
    client = await db.clients.find_one({"id": followup['client_id']}, {"_id": 0})
    
    prompt = f"""Generate a professional follow-up email for Anka Healthcare CS team.

Client: {followup['client_name']}
Issue: {followup['description']}
Suggested Action: {followup['suggested_action']}
Days Open: {followup.get('days_open', 0)}

Write a warm but professional email that:
- Acknowledges the issue
- Proposes next steps
- Maintains healthcare industry professionalism
- Includes a clear call-to-action

Format as JSON:
{{"subject": "...", "body": "..."}}"""
    
    try:
        response = await generate_with_ai(prompt)
        import json
        email_data = json.loads(response)
    except:
        email_data = {
            "subject": f"Follow-up: {followup['description'][:50]}",
            "body": f"Dear Team,\n\nI wanted to follow up regarding {followup['description']}.\n\n{followup['suggested_action']}\n\nPlease let me know if you have any questions.\n\nBest regards,\nAnka Healthcare CS Team"
        }
    
    # Create draft communication
    comm = Communication(
        client_id=followup['client_id'],
        comm_type="draft",
        channel="email",
        subject=email_data['subject'],
        body=email_data['body'],
        ai_generated=True,
        template_used="followup_draft",
        created_by=user['id']
    )
    await db.communications.insert_one(comm.model_dump())
    
    # Link draft to followup
    await db.followups.update_one({"id": followup_id}, {"$set": {"draft_email_id": comm.id}})
    
    return {"draft_id": comm.id, "subject": email_data['subject'], "body": email_data['body']}

# =============================================================================
# AI PROMPT TEMPLATES
# =============================================================================

@api_router.get("/templates", response_model=List[PromptTemplate])
async def get_templates(user: dict = Depends(require_auth)):
    templates = await db.prompt_templates.find({}, {"_id": 0}).to_list(100)
    return templates

@api_router.post("/templates/execute/{slug}")
async def execute_template(slug: str, params: Dict[str, Any], user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Execute an AI template with given parameters"""
    template = await db.prompt_templates.find_one({"slug": slug}, {"_id": 0})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Get client data if client_id provided
    client_data = {}
    if params.get('client_id'):
        client = await db.clients.find_one({"id": params['client_id']}, {"_id": 0})
        if client:
            client_data = client
            
            # Get recent performance
            perf = await db.performance_records.find({"client_id": params['client_id']}).sort("period_end", -1).limit(3).to_list(3)
            client_data['recent_performance'] = perf
            
            # Get contacts
            contacts = await db.client_contacts.find({"client_id": params['client_id']}, {"_id": 0}).to_list(100)
            client_data['contacts'] = contacts
            
            # Get recent communications
            comms = await db.communications.find({"client_id": params['client_id']}).sort("created_at", -1).limit(10).to_list(10)
            client_data['recent_communications'] = comms
    
    # Build prompt from template
    prompt_vars = {**params, "client_data": client_data}
    user_prompt = template['user_prompt_template'].format(**prompt_vars)
    
    # Scrub PHI
    scrubbed_prompt, redactions = scrub_phi(user_prompt, f"template_{slug}")
    await log_phi_redactions(redactions)
    
    # Generate
    response = await generate_with_ai(scrubbed_prompt, template['system_prompt'])
    
    # Audit log
    await log_audit(user['id'], user['name'], 'ai_generation', 'template', slug, {'client_id': params.get('client_id')})
    
    return {"template": slug, "result": response}

# =============================================================================
# CRM SYNC ROUTES (HubSpot-like structure)
# =============================================================================

@api_router.get("/crm/contacts", response_model=List[CRMContact])
async def get_crm_contacts(user: dict = Depends(require_auth)):
    contacts = await db.crm_contacts.find({}, {"_id": 0}).to_list(1000)
    return contacts

@api_router.post("/crm/contacts", response_model=CRMContact)
async def create_crm_contact(contact: CRMContact, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    await db.crm_contacts.insert_one(contact.model_dump())
    return contact

@api_router.get("/crm/deals", response_model=List[CRMDeal])
async def get_crm_deals(user: dict = Depends(require_auth)):
    deals = await db.crm_deals.find({}, {"_id": 0}).to_list(1000)
    return deals

@api_router.post("/crm/deals", response_model=CRMDeal)
async def create_crm_deal(deal: CRMDeal, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    await db.crm_deals.insert_one(deal.model_dump())
    return deal

@api_router.get("/crm/tickets", response_model=List[CRMTicket])
async def get_crm_tickets(user: dict = Depends(require_auth)):
    tickets = await db.crm_tickets.find({}, {"_id": 0}).to_list(1000)
    return tickets

@api_router.post("/crm/sync")
async def trigger_crm_sync(user: dict = Depends(require_role(["cs_lead"]))):
    """Trigger manual CRM sync - stub for future HubSpot integration"""
    # This would connect to HubSpot API when API key is configured
    integration = await db.integrations.find_one({"integration_name": "hubspot"}, {"_id": 0})
    
    if not integration or integration.get('connection_status') != 'connected':
        return {"status": "skipped", "message": "HubSpot integration not configured"}
    
    # Placeholder for actual sync logic
    return {"status": "completed", "message": "CRM sync completed", "synced_at": datetime.now(timezone.utc).isoformat()}

# =============================================================================
# ALERTING ENGINE
# =============================================================================

@api_router.post("/alerts/check-engagement")
async def check_engagement_gaps(user: dict = Depends(require_role(["cs_lead"]))):
    """Check for engagement gaps and create alerts"""
    alerts_created = 0
    
    # Get thresholds
    default_threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
    if not default_threshold:
        default_threshold = {"engagement_gap_decision_maker_days": 14, "engagement_gap_influencer_days": 30}
    
    contacts = await db.client_contacts.find({}, {"_id": 0}).to_list(10000)
    now = datetime.now(timezone.utc)
    
    for contact in contacts:
        if contact['role_type'] not in ['decision-maker', 'influencer']:
            continue
        
        # Get client-specific threshold or use default
        threshold = await db.alert_thresholds.find_one({"client_id": contact['client_id']}, {"_id": 0})
        if not threshold:
            threshold = default_threshold
        
        days_threshold = (
            threshold['engagement_gap_decision_maker_days']
            if contact['role_type'] == 'decision-maker'
            else threshold['engagement_gap_influencer_days']
        )
        
        # Check all engagement dates
        last_dates = []
        for field in ['last_email_date', 'last_call_date', 'last_meeting_date']:
            if contact.get(field):
                last_dates.append(datetime.fromisoformat(contact[field].replace('Z', '+00:00')))
        
        if not last_dates:
            # No engagement ever - create alert
            most_recent = None
            days_since = 999
        else:
            most_recent = max(last_dates)
            days_since = (now - most_recent).days
        
        if days_since > days_threshold:
            # Check if alert already exists
            existing = await db.alerts.find_one({
                "client_id": contact['client_id'],
                "alert_type": "engagement_gap",
                "status": {"$in": ["active", "acknowledged"]},
                "trigger_data.contact_id": contact['id']
            }, {"_id": 0})
            
            if not existing:
                client = await db.clients.find_one({"id": contact['client_id']}, {"_id": 0})
                client_name = client.get('name', 'Unknown') if client else 'Unknown'
                
                severity = "high" if contact['role_type'] == 'decision-maker' else "medium"
                
                alert = Alert(
                    client_id=contact['client_id'],
                    client_name=client_name,
                    alert_type="engagement_gap",
                    severity=severity,
                    title=f"No contact with {contact['name']} ({contact['role_type']}) in {days_since} days",
                    description=f"Contact {contact['name']} ({contact['title']}) has not been engaged via email, call, or meeting in {days_since} days.",
                    trigger_data={
                        "contact_id": contact['id'],
                        "contact_name": contact['name'],
                        "role_type": contact['role_type'],
                        "days_since_contact": days_since,
                        "threshold_days": days_threshold
                    }
                )
                await db.alerts.insert_one(alert.model_dump())
                alerts_created += 1
    
    return {"message": f"Created {alerts_created} engagement gap alerts"}

@api_router.post("/alerts/check-renewals")
async def check_renewal_alerts(user: dict = Depends(require_role(["cs_lead"]))):
    """Check for upcoming renewals and create alerts"""
    alerts_created = 0
    now = datetime.now(timezone.utc)
    
    clients = await db.clients.find({"status": "active"}, {"_id": 0}).to_list(10000)
    
    for client in clients:
        if not client.get('contract_end'):
            continue
        
        contract_end_str = client['contract_end']
        # Handle various date formats
        if 'T' in contract_end_str or '+' in contract_end_str or 'Z' in contract_end_str:
            contract_end = datetime.fromisoformat(contract_end_str.replace('Z', '+00:00'))
        else:
            # Simple date string like "2024-01-01"
            contract_end = datetime.fromisoformat(contract_end_str + "T00:00:00+00:00")
        days_until = (contract_end - now).days
        
        # Get threshold settings
        threshold = await db.alert_thresholds.find_one({"client_id": client['id']}, {"_id": 0})
        if not threshold:
            threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
        if not threshold:
            threshold = {"renewal_alert_60_days": True, "renewal_alert_30_days": True, "renewal_alert_15_days": True}
        
        alert_days = []
        if threshold.get('renewal_alert_60_days') and 55 <= days_until <= 65:
            alert_days.append((60, "medium"))
        if threshold.get('renewal_alert_30_days') and 25 <= days_until <= 35:
            alert_days.append((30, "high"))
        if threshold.get('renewal_alert_15_days') and 10 <= days_until <= 20:
            alert_days.append((15, "critical"))
        
        for days, severity in alert_days:
            # Check if alert already exists
            existing = await db.alerts.find_one({
                "client_id": client['id'],
                "alert_type": "renewal",
                "status": {"$in": ["active", "acknowledged"]},
                "trigger_data.days_threshold": days
            }, {"_id": 0})
            
            if not existing:
                alert = Alert(
                    client_id=client['id'],
                    client_name=client['name'],
                    alert_type="renewal",
                    severity=severity,
                    title=f"Contract renewal in {days_until} days",
                    description=f"Contract for {client['name']} expires on {client['contract_end']}. Schedule QBR if not already done.",
                    trigger_data={
                        "contract_end": client['contract_end'],
                        "days_until_renewal": days_until,
                        "days_threshold": days
                    }
                )
                await db.alerts.insert_one(alert.model_dump())
                alerts_created += 1
    
    return {"message": f"Created {alerts_created} renewal alerts"}

@api_router.post("/alerts/check-overdue-followups")
async def check_overdue_followups(user: dict = Depends(require_role(["cs_lead"]))):
    """Check for overdue follow-ups and create alerts"""
    alerts_created = 0
    now = datetime.now(timezone.utc)
    
    # Get threshold
    threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
    overdue_days = threshold.get('followup_overdue_days', 7) if threshold else 7
    
    followups = await db.followups.find({"status": "open"}, {"_id": 0}).to_list(10000)
    
    for followup in followups:
        created = datetime.fromisoformat(followup['created_at'].replace('Z', '+00:00'))
        days_open = (now - created).days
        
        if days_open > overdue_days:
            # Check if alert already exists
            existing = await db.alerts.find_one({
                "alert_type": "followup_overdue",
                "status": {"$in": ["active", "acknowledged"]},
                "trigger_data.followup_id": followup['id']
            }, {"_id": 0})
            
            if not existing:
                alert = Alert(
                    client_id=followup['client_id'],
                    client_name=followup['client_name'],
                    alert_type="followup_overdue",
                    severity="medium",
                    title=f"Follow-up overdue by {days_open - overdue_days} days",
                    description=f"Follow-up '{followup['description'][:100]}' has been open for {days_open} days.",
                    trigger_data={
                        "followup_id": followup['id'],
                        "days_open": days_open,
                        "threshold_days": overdue_days,
                        "owner": followup['owner']
                    },
                    assigned_to=followup['owner']
                )
                await db.alerts.insert_one(alert.model_dump())
                alerts_created += 1
    
    return {"message": f"Created {alerts_created} overdue follow-up alerts"}

# =============================================================================
# ALERT THRESHOLDS (Settings)
# =============================================================================

@api_router.get("/settings/alert-thresholds")
async def get_alert_thresholds(client_id: Optional[str] = None, user: dict = Depends(require_role(["cs_lead"]))):
    query = {"client_id": client_id} if client_id else {"client_id": None}
    threshold = await db.alert_thresholds.find_one(query, {"_id": 0})
    
    if not threshold:
        # Return defaults
        threshold = AlertThreshold(client_id=client_id).model_dump()
    
    return threshold

@api_router.post("/settings/alert-thresholds")
async def set_alert_thresholds(threshold: AlertThreshold, user: dict = Depends(require_role(["cs_lead"]))):
    # Upsert threshold
    await db.alert_thresholds.update_one(
        {"client_id": threshold.client_id},
        {"$set": threshold.model_dump()},
        upsert=True
    )
    return {"message": "Thresholds updated"}

# =============================================================================
# DASHBOARD METRICS
# =============================================================================

@api_router.get("/dashboard/metrics")
async def get_dashboard_metrics(user: dict = Depends(require_auth)):
    """Get key metrics for dashboard"""
    query = {}
    if user['role'] == 'csm':
        query["id"] = {"$in": user.get('assigned_clients', [])}
    
    # Total clients
    total_clients = await db.clients.count_documents(query)
    
    # At-risk clients
    at_risk_query = {**query, "status": "at-risk"}
    at_risk = await db.clients.count_documents(at_risk_query)
    
    # Average recovery rate (last period)
    pipeline = [
        {"$sort": {"period_end": -1}},
        {"$group": {"_id": "$client_id", "latest_rate": {"$first": "$recovery_rate"}}},
        {"$group": {"_id": None, "avg_rate": {"$avg": "$latest_rate"}}}
    ]
    result = await db.performance_records.aggregate(pipeline).to_list(1)
    avg_recovery_rate = result[0]['avg_rate'] if result else 0
    
    # Open follow-ups
    followup_query = {"status": "open"}
    if user['role'] == 'csm':
        followup_query["owner"] = user['id']
    open_followups = await db.followups.count_documents(followup_query)
    
    # Active alerts
    alert_query = {"status": {"$in": ["active", "acknowledged"]}}
    if user['role'] == 'csm':
        alert_query["client_id"] = {"$in": user.get('assigned_clients', [])}
    active_alerts = await db.alerts.count_documents(alert_query)
    
    return {
        "total_clients": total_clients,
        "at_risk_clients": at_risk,
        "avg_recovery_rate": round(avg_recovery_rate, 1),
        "open_followups": open_followups,
        "active_alerts": active_alerts
    }

# =============================================================================
# AUDIT LOGS
# =============================================================================

@api_router.get("/audit-logs")
async def get_audit_logs(
    user_id: Optional[str] = None,
    action_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = 100,
    admin: dict = Depends(require_role(["cs_lead"]))
):
    """Get audit logs with filters"""
    query = {}
    if user_id:
        query["user_id"] = user_id
    if action_type:
        query["action_type"] = action_type
    if start_date:
        query["timestamp"] = {"$gte": start_date}
    if end_date:
        if "timestamp" in query:
            query["timestamp"]["$lte"] = end_date
        else:
            query["timestamp"] = {"$lte": end_date}
    
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("timestamp", -1).limit(limit).to_list(limit)
    return logs

# =============================================================================
# INNOVATION BRIEFING
# =============================================================================

@api_router.post("/clients/{client_id}/mark-innovation-briefed")
async def mark_innovation_briefed(client_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Mark a client as having received an innovation briefing"""
    result = await db.clients.update_one(
        {"id": client_id},
        {"$set": {"last_innovation_briefing": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    
    await log_audit(user['id'], user['name'], 'alert_action', 'client', client_id, {'action': 'innovation_briefed'})
    return {"message": "Client marked as innovation briefed"}

# =============================================================================
# POLICY UPDATES
# =============================================================================

@api_router.get("/policy-updates")
async def get_policy_updates(user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Get all policy updates"""
    updates = await db.policy_updates.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return updates

@api_router.post("/policy-updates")
async def create_policy_update(input: PolicyUpdate, user: dict = Depends(require_role(["cs_lead"]))):
    """Create a new policy update"""
    input.created_by = user['id']
    await db.policy_updates.insert_one(input.model_dump())
    return input

@api_router.patch("/policy-updates/{policy_id}")
async def update_policy(policy_id: str, updates: Dict[str, Any], user: dict = Depends(require_role(["cs_lead"]))):
    """Update a policy update"""
    result = await db.policy_updates.update_one({"id": policy_id}, {"$set": updates})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Policy update not found")
    return {"message": "Policy updated"}

@api_router.delete("/policy-updates/{policy_id}")
async def delete_policy(policy_id: str, user: dict = Depends(require_role(["cs_lead"]))):
    """Delete a policy update"""
    result = await db.policy_updates.delete_one({"id": policy_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Policy update not found")
    return {"message": "Policy deleted"}

@api_router.post("/policy-updates/{policy_id}/generate-alerts")
async def generate_policy_alerts(policy_id: str, user: dict = Depends(require_role(["cs_lead"]))):
    """Generate alerts for all clients affected by a policy update"""
    policy = await db.policy_updates.find_one({"id": policy_id}, {"_id": 0})
    if not policy:
        raise HTTPException(status_code=404, detail="Policy update not found")
    
    clients = await db.clients.find({"status": "active"}, {"_id": 0}).to_list(1000)
    alerts_created = 0
    
    for client in clients:
        # Check for service overlap
        client_services = set(client.get('contracted_services', []))
        policy_services = set(policy.get('affected_services', []))
        overlapping = client_services.intersection(policy_services)
        
        if overlapping:
            # Check for existing alert
            existing = await db.alerts.find_one({
                "client_id": client['id'],
                "alert_type": "policy_update",
                "trigger_data.policy_id": policy_id,
                "status": {"$in": ["active", "acknowledged"]}
            }, {"_id": 0})
            
            if not existing:
                alert = Alert(
                    client_id=client['id'],
                    client_name=client['name'],
                    alert_type="policy_update",
                    severity="medium",
                    title=f"Policy update: {policy['title']} affects {client['name']}",
                    description=policy['description'],
                    trigger_data={"policy_id": policy_id, "matching_services": list(overlapping)}
                )
                await db.alerts.insert_one(alert.model_dump())
                alerts_created += 1
    
    return {"message": f"Generated {alerts_created} alerts", "alerts_created": alerts_created}

# =============================================================================
# EMAIL CADENCES
# =============================================================================

@api_router.get("/email-cadences")
async def get_email_cadences(client_id: Optional[str] = None, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Get email cadences, optionally filtered by client"""
    query = {}
    if client_id:
        query["client_id"] = client_id
    cadences = await db.email_cadences.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return cadences

@api_router.post("/email-cadences")
async def create_email_cadence(input: EmailCadence, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Create a new email cadence"""
    input.created_by = user['id']
    input.next_trigger = (datetime.now(timezone.utc) + timedelta(days=input.frequency_days)).isoformat()
    await db.email_cadences.insert_one(input.model_dump())
    return input

@api_router.patch("/email-cadences/{cadence_id}")
async def update_email_cadence(cadence_id: str, status: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Update email cadence status"""
    result = await db.email_cadences.update_one({"id": cadence_id}, {"$set": {"status": status}})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Cadence not found")
    return {"message": "Cadence updated"}

@api_router.delete("/email-cadences/{cadence_id}")
async def delete_email_cadence(cadence_id: str, user: dict = Depends(require_role(["cs_lead"]))):
    """Delete an email cadence"""
    result = await db.email_cadences.delete_one({"id": cadence_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cadence not found")
    return {"message": "Cadence deleted"}

@api_router.post("/email-cadences/{cadence_id}/send")
async def trigger_email_cadence(cadence_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Manually trigger an email cadence"""
    cadence = await db.email_cadences.find_one({"id": cadence_id}, {"_id": 0})
    if not cadence:
        raise HTTPException(status_code=404, detail="Cadence not found")
    
    # Get client data
    client = await db.clients.find_one({"id": cadence['client_id']}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Update cadence timestamps
    now = datetime.now(timezone.utc)
    await db.email_cadences.update_one({"id": cadence_id}, {"$set": {
        "last_triggered": now.isoformat(),
        "next_trigger": (now + timedelta(days=cadence['frequency_days'])).isoformat()
    }})
    
    return {"message": "Cadence triggered", "client": client['name']}

# =============================================================================
# WEEKLY SUMMARY ENDPOINT
# =============================================================================

@api_router.post("/weekly-summary")
async def generate_weekly_summary(data: Dict[str, Any], user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Generate AI-powered weekly summary for a client"""
    client_id = data.get('client_id')
    week_ending = data.get('week_ending')
    
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Gather context
    performance = await db.performance_records.find({"client_id": client_id}).sort("period_end", -1).limit(3).to_list(3)
    comms = await db.communications.find({"client_id": client_id}).sort("created_at", -1).limit(20).to_list(20)
    alerts = await db.alerts.find({"client_id": client_id, "status": {"$in": ["active", "acknowledged"]}}).to_list(50)
    followups = await db.followups.find({"client_id": client_id, "status": "open"}).to_list(50)
    
    prompt = f'''Generate a weekly client summary for {client['name']} for the week ending {week_ending}.

Performance Data: {json.dumps([{k:v for k,v in p.items() if k != '_id'} for p in performance], default=str)}
Recent Communications ({len(comms)} this period): {json.dumps([{'subject': c.get('subject'), 'channel': c.get('channel'), 'date': c.get('created_at')} for c in comms], default=str)}
Active Alerts: {json.dumps([{'title': a.get('title'), 'severity': a.get('severity'), 'type': a.get('alert_type')} for a in alerts], default=str)}
Open Follow-ups: {json.dumps([{'title': f.get('title'), 'priority_level': f.get('priority_level'), 'days_open': f.get('days_open', 0)} for f in followups], default=str)}

Respond ONLY with valid JSON (no markdown, no backticks):
{{"key_topics": ["topic1", "topic2"],
 "action_items": [{{"item": "description", "owner": "name", "priority": "high/medium/low"}}],
 "red_flags": ["flag1", "flag2"],
 "new_contacts": ["contact info if detected"],
 "escalation_items": []}}'''
    
    scrubbed, _ = scrub_phi(prompt, 'weekly_summary')
    response = await generate_with_ai(scrubbed)
    
    # Parse JSON response
    try:
        clean = response.replace('```json', '').replace('```', '').strip()
        result = json.loads(clean)
    except:
        result = {'key_topics': ['Summary generation completed'], 'action_items': [], 'red_flags': [], 'new_contacts': [], 'escalation_items': []}
    
    await log_audit(user['id'], user['name'], 'ai_generation', 'weekly_summary', client_id, {'week_ending': week_ending})
    return result

# =============================================================================
# SCHEDULED REVIEWS (Calendar)
# =============================================================================

@api_router.get("/reviews")
async def get_reviews(
    client_id: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    user: dict = Depends(require_auth)
):
    query = {}
    if client_id:
        query["client_id"] = client_id
    if start_date:
        query["scheduled_date"] = {"$gte": start_date}
    if end_date:
        if "scheduled_date" in query:
            query["scheduled_date"]["$lte"] = end_date
        else:
            query["scheduled_date"] = {"$lte": end_date}
    
    reviews = await db.scheduled_reviews.find(query, {"_id": 0}).sort("scheduled_date", 1).to_list(1000)
    return reviews

@api_router.post("/reviews")
async def create_review(
    client_id: str,
    client_name: str,
    review_type: str,
    scheduled_date: str,
    scheduled_time: str,
    duration_minutes: int = 60,
    attendees: List[str] = [],
    notes: Optional[str] = None,
    user: dict = Depends(require_role(["cs_lead", "csm"]))
):
    review = ScheduledReview(
        client_id=client_id,
        client_name=client_name,
        review_type=review_type,
        scheduled_date=scheduled_date,
        scheduled_time=scheduled_time,
        duration_minutes=duration_minutes,
        attendees=attendees,
        notes=notes
    )
    await db.scheduled_reviews.insert_one(review.model_dump())
    return review

@api_router.patch("/reviews/{review_id}")
async def update_review(review_id: str, status: Optional[str] = None, notes: Optional[str] = None, user: dict = Depends(require_auth)):
    update_data = {}
    if status:
        update_data["status"] = status
    if notes is not None:
        update_data["notes"] = notes
    
    if update_data:
        result = await db.scheduled_reviews.update_one({"id": review_id}, {"$set": update_data})
        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="Review not found")
    
    return {"message": "Review updated"}

@api_router.delete("/reviews/{review_id}")
async def delete_review(review_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    result = await db.scheduled_reviews.delete_one({"id": review_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Review not found")
    return {"message": "Review deleted"}

# =============================================================================
# TEAM MEMBERS (Org Mapping)
# =============================================================================

@api_router.get("/team-members")
async def get_team_members(user: dict = Depends(require_auth)):
    members = await db.team_members.find({}, {"_id": 0}).to_list(1000)
    return members

@api_router.post("/team-members")
async def create_team_member(name: str, role: str, email: str, phone: Optional[str] = None, user: dict = Depends(require_role(["cs_lead"]))):
    member = TeamMember(name=name, role=role, email=email, phone=phone)
    await db.team_members.insert_one(member.model_dump())
    return member

@api_router.delete("/team-members/{member_id}")
async def delete_team_member(member_id: str, user: dict = Depends(require_role(["cs_lead"]))):
    result = await db.team_members.delete_one({"id": member_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Team member not found")
    return {"message": "Team member deleted"}

# =============================================================================
# LOBs (Lines of Business)
# =============================================================================

@api_router.get("/lobs")
async def get_lobs(user: dict = Depends(require_auth)):
    lobs = await db.lobs.find({}, {"_id": 0}).to_list(1000)
    return lobs

@api_router.post("/lobs")
async def create_lob(name: str, description: Optional[str] = None, user: dict = Depends(require_role(["cs_lead"]))):
    lob = LOB(name=name, description=description)
    await db.lobs.insert_one(lob.model_dump())
    return lob

# =============================================================================
# CLIENT LOB MAPPINGS
# =============================================================================

@api_router.get("/org-mappings")
async def get_org_mappings(client_id: Optional[str] = None, user: dict = Depends(require_auth)):
    query = {}
    if client_id:
        query["client_id"] = client_id
    mappings = await db.client_lob_mappings.find(query, {"_id": 0}).to_list(1000)
    return mappings

@api_router.post("/org-mappings")
async def create_org_mapping(
    client_id: str,
    client_name: str,
    lob_id: str,
    lob_name: str,
    anka_team_members: List[str] = [],
    user: dict = Depends(require_role(["cs_lead", "csm"]))
):
    mapping = ClientLOBMapping(
        client_id=client_id,
        client_name=client_name,
        lob_id=lob_id,
        lob_name=lob_name,
        anka_team_members=anka_team_members
    )
    await db.client_lob_mappings.insert_one(mapping.model_dump())
    return mapping

@api_router.delete("/org-mappings/{mapping_id}")
async def delete_org_mapping(mapping_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    result = await db.client_lob_mappings.delete_one({"id": mapping_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Mapping not found")
    return {"message": "Mapping deleted"}

# =============================================================================
# INTEGRATIONS CONFIG
# =============================================================================

@api_router.get("/integrations")
async def get_integrations(user: dict = Depends(require_role(["cs_lead"]))):
    integrations = await db.integrations.find({}, {"_id": 0}).to_list(100)
    return integrations

@api_router.post("/integrations")
async def create_integration(integration_name: str, config: Dict[str, Any] = {}, user: dict = Depends(require_role(["cs_lead"]))):
    integration = IntegrationConfig(integration_name=integration_name, config=config)
    await db.integrations.update_one(
        {"integration_name": integration_name},
        {"$set": integration.model_dump()},
        upsert=True
    )
    return integration

@api_router.patch("/integrations/{integration_name}/status")
async def update_integration_status(integration_name: str, status: str, user: dict = Depends(require_role(["cs_lead"]))):
    await db.integrations.update_one(
        {"integration_name": integration_name},
        {"$set": {"connection_status": status, "last_sync": datetime.now(timezone.utc).isoformat()}}
    )
    return {"message": "Integration status updated"}

# =============================================================================
# REPORT GENERATION (DOCX)
# =============================================================================

@api_router.post("/reports/generate-docx")
async def generate_docx_report(client_id: str, report_type: str = "external", user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Generate a Word document report for a client"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Fetch client data
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Fetch performance data
    performance = await db.performance_records.find({"client_id": client_id}).sort("period_end", -1).limit(3).to_list(3)
    
    # Generate AI narrative
    try:
        template_slug = "client_report_external" if report_type == "external" else "weekly_summary"
        ai_prompt = f"""Generate a professional {report_type} report narrative for {client['name']}.

Performance Data:
{json.dumps(performance, indent=2, default=str)}

Client Status: {client.get('status', 'active')}
Health Score: {client.get('health_score', 80)}

Write a concise executive summary (2-3 paragraphs) suitable for a {'client-facing' if report_type == 'external' else 'internal team'} document."""
        
        ai_narrative = await generate_with_ai(ai_prompt)
    except Exception:
        ai_narrative = f"Performance report for {client['name']}. Please review the metrics below for detailed analysis."
    
    # Create document
    doc = Document()
    
    # Header
    header = doc.add_heading('Anka Healthcare', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title = doc.add_heading(f"{'Client Performance Report' if report_type == 'external' else 'Internal Performance Report'}", 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Client info
    doc.add_paragraph(f"Client: {client['name']}")
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Report Type: {report_type.title()}")
    doc.add_paragraph("")
    
    # Executive Summary
    doc.add_heading('Executive Summary', 2)
    doc.add_paragraph(ai_narrative)
    
    # Performance Metrics
    doc.add_heading('Performance Metrics', 2)
    
    if performance:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Period'
        hdr_cells[1].text = 'Denials Worked'
        hdr_cells[2].text = 'Recovered'
        hdr_cells[3].text = 'Recovery Rate'
        hdr_cells[4].text = 'SLA Compliance'
        
        for perf in performance:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{perf.get('period_start', 'N/A')[:10]} - {perf.get('period_end', 'N/A')[:10]}"
            row_cells[1].text = f"{perf.get('denials_worked', 0):,}"
            row_cells[2].text = f"${perf.get('dollars_recovered', 0):,.2f}"
            row_cells[3].text = f"{perf.get('recovery_rate', 0):.1f}%"
            row_cells[4].text = f"{perf.get('sla_compliance_pct', 0):.1f}%"
    else:
        doc.add_paragraph("No performance data available.")
    
    # Internal-only sections
    if report_type == "internal":
        doc.add_heading('Internal Notes', 2)
        doc.add_paragraph(f"Client Status: {client.get('status', 'active')}")
        doc.add_paragraph(f"Health Score: {client.get('health_score', 80)}")
        doc.add_paragraph(f"Assigned CSM: {client.get('assigned_csm', 'N/A')}")
        
        # Recent alerts
        alerts = await db.alerts.find({"client_id": client_id, "status": {"$in": ["active", "acknowledged"]}}).limit(5).to_list(5)
        if alerts:
            doc.add_heading('Active Alerts', 3)
            for alert in alerts:
                doc.add_paragraph(f"• [{alert.get('severity', 'medium').upper()}] {alert.get('title', 'N/A')}")
    
    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("Generated by Anka CS Hub")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Log audit
    await log_audit(user['id'], user['name'], 'report_download', 'report', client_id, {
        'client_name': client['name'],
        'report_type': report_type
    })
    
    filename = f"anka_report_{client['name'].replace(' ', '_')}_{report_type}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.post("/reports/generate-both")
async def generate_both_reports(client_id: str, user: dict = Depends(require_role(["cs_lead", "csm"]))):
    """Generate both internal and external reports - returns URLs to download each"""
    # For simplicity, we'll return instructions to call the endpoint twice
    return {
        "message": "Call /reports/generate-docx twice with report_type=internal and report_type=external",
        "internal_url": f"/api/reports/generate-docx?client_id={client_id}&report_type=internal",
        "external_url": f"/api/reports/generate-docx?client_id={client_id}&report_type=external"
    }

@api_router.get("/")
async def root():
    return {"message": "Anka CS Hub API v2.0", "status": "healthy"}

@api_router.get("/health")
async def health():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# =============================================================================
# APP SETUP
# =============================================================================

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# =============================================================================
# STARTUP: Seed initial data and templates
# =============================================================================

@app.on_event("startup")
async def startup_event():
    # Create indexes
    await db.users.create_index("email", unique=True)
    await db.clients.create_index("name")
    await db.alerts.create_index([("client_id", 1), ("status", 1)])
    await db.performance_records.create_index([("client_id", 1), ("period_end", -1)])
    
    # Seed default admin user if none exists
    admin = await db.users.find_one({"role": "cs_lead"}, {"_id": 0})
    if not admin:
        admin_user = User(
            name="Admin User",
            email="admin@anka.health",
            password_hash=hash_password("admin123"),
            role="cs_lead"
        )
        await db.users.insert_one(admin_user.model_dump())
        logger.info("Created default admin user: admin@anka.health / admin123")
    
    # Seed prompt templates if none exist
    templates_count = await db.prompt_templates.count_documents({})
    if templates_count == 0:
        templates = [
            PromptTemplate(
                name="Weekly Client Summary",
                slug="weekly_summary",
                description="Aggregates all communications and performance data for a client in a given week",
                system_prompt="You are an expert healthcare customer success analyst. Generate comprehensive weekly summaries that highlight key issues, wins, and action items.",
                user_prompt_template="""Generate a weekly summary for {client_data[name]} for the week ending {date_range}.

Client Data:
- Status: {client_data[status]}
- Health Score: {client_data[health_score]}
- Recent Performance: {client_data[recent_performance]}
- Recent Communications: {client_data[recent_communications]}

Provide a structured summary with:
1. Key Topics Discussed
2. Action Items with Owners
3. Red Flags or Concerns
4. New Contacts Detected
5. Items Requiring Escalation

Format as JSON.""",
                parameters=["client_id", "date_range"]
            ),
            PromptTemplate(
                name="Client Performance Report (External)",
                slug="client_report_external",
                description="Professional client-facing performance report",
                system_prompt="You are a healthcare revenue cycle reporting specialist. Generate professional, client-facing reports that highlight performance metrics and improvement areas.",
                user_prompt_template="""Generate an external performance report for {client_data[name]}.

Performance Data:
{client_data[recent_performance]}

Create a professional report with:
1. Executive Summary
2. Denials Worked & Recovered
3. Recovery Rate vs SLA Target
4. Top Denial Codes Analysis
5. Payer Breakdown
6. Trend vs Prior Period
7. Next Week Focus Areas (AI recommendations)

Keep tone professional and positive while being transparent about areas for improvement.""",
                parameters=["client_id", "date_range"]
            ),
            PromptTemplate(
                name="Follow-Up List Generator",
                slug="followup_list",
                description="Generates prioritized follow-up list with suggested actions",
                system_prompt="You are a customer success operations specialist. Generate actionable follow-up lists with clear priorities and suggested approaches.",
                user_prompt_template="""Generate a prioritized follow-up list based on:

Open Items: {open_items}
Client Health Scores: {health_scores}
Upcoming Renewals: {renewals}

For each item provide:
1. Priority Score (1-10)
2. Category (Call Required / Email Sufficient)
3. Suggested Action
4. Draft talking points or email snippet""",
                parameters=["open_items", "health_scores", "renewals"]
            ),
            PromptTemplate(
                name="Engagement Gap Outreach",
                slug="engagement_outreach",
                description="Warm check-in email for stale contacts",
                system_prompt="You are a relationship-focused customer success manager. Write warm, personalized outreach emails that re-engage contacts without being pushy.",
                user_prompt_template="""Generate a warm check-in email for {contact_name} at {client_data[name]}.

Contact Role: {contact_role}
Days Since Last Contact: {days_since_contact}
Recent Client Performance: {client_data[recent_performance]}
Recent Issues: {recent_issues}

Write an email that:
- Feels personal, not templated
- References specific recent work or results
- Offers value (insight, update, resource)
- Has a soft call-to-action

Format as JSON with subject and body.""",
                parameters=["client_id", "contact_name", "contact_role", "days_since_contact", "recent_issues"]
            ),
            PromptTemplate(
                name="QBR Scheduling",
                slug="qbr_scheduling",
                description="QBR scheduling email with proposed agenda",
                system_prompt="You are an executive assistant for healthcare customer success. Write professional meeting scheduling emails with clear agendas.",
                user_prompt_template="""Generate a QBR scheduling email for {client_data[name]}.

Contract End Date: {client_data[contract_end]}
Recent Performance Summary: {client_data[recent_performance]}
Key Contacts: {client_data[contacts]}

Create an email that:
- Proposes 2-3 potential meeting times
- Includes a draft agenda based on performance data
- Lists recommended attendees
- Sets expectations for the meeting

Format as JSON with subject, body, and proposed_agenda.""",
                parameters=["client_id"]
            ),
            PromptTemplate(
                name="Escalation Notification",
                slug="escalation_notification",
                description="Internal escalation email with context and recommendations",
                system_prompt="You are a senior customer success leader. Write clear, action-oriented internal escalation notifications.",
                user_prompt_template="""Generate an internal escalation notification for {client_data[name]}.

Escalation Trigger: {trigger_reason}
Alert Details: {alert_details}
Client History: {client_data[recent_communications]}
Current Health Score: {client_data[health_score]}

Create an internal email that:
- Clearly states the issue and its severity
- Provides relevant context
- Recommends specific actions
- Identifies who should be involved

Format as JSON with subject, body, recommended_actions, and suggested_attendees.""",
                parameters=["client_id", "trigger_reason", "alert_details"]
            )
        ]
        
        for template in templates:
            await db.prompt_templates.insert_one(template.model_dump())
        logger.info(f"Seeded {len(templates)} prompt templates")
    
    # Seed default alert thresholds if none exist
    default_threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
    if not default_threshold:
        threshold = AlertThreshold(client_id=None)
        await db.alert_thresholds.insert_one(threshold.model_dump())
        logger.info("Created default alert thresholds")
    
    # Seed default LOBs if none exist
    lobs_count = await db.lobs.count_documents({})
    if lobs_count == 0:
        default_lobs = [
            LOB(name="EV", description="Eligibility Verification"),
            LOB(name="Prior Auth", description="Prior Authorization"),
            LOB(name="Coding", description="Medical Coding"),
            LOB(name="Billing", description="Medical Billing"),
            LOB(name="AR", description="Accounts Receivable"),
            LOB(name="Payment Posting", description="Payment Posting & Reconciliation"),
            LOB(name="Denial Management", description="Denial Appeals & Management")
        ]
        for lob in default_lobs:
            await db.lobs.insert_one(lob.model_dump())
        logger.info(f"Seeded {len(default_lobs)} LOBs")
    
    # Seed default team members if none exist
    members_count = await db.team_members.count_documents({})
    if members_count == 0:
        default_members = [
            TeamMember(name="Sarah Mitchell", role="manager", email="sarah.mitchell@anka.health"),
            TeamMember(name="James Rodriguez", role="supervisor", email="james.rodriguez@anka.health"),
            TeamMember(name="Emily Chen", role="team_lead", email="emily.chen@anka.health"),
            TeamMember(name="Michael Thompson", role="supervisor", email="michael.thompson@anka.health")
        ]
        for member in default_members:
            await db.team_members.insert_one(member.model_dump())
        logger.info(f"Seeded {len(default_members)} team members")
    
    # Seed default integrations if none exist
    integrations_count = await db.integrations.count_documents({})
    if integrations_count == 0:
        default_integrations = [
            IntegrationConfig(integration_name="microsoft_365", connection_status="disconnected", config={"tenant_id": "", "client_id": ""}),
            IntegrationConfig(integration_name="hubspot", connection_status="disconnected", config={"api_key": ""}),
            IntegrationConfig(integration_name="google_calendar", connection_status="disconnected", config={"credentials": {}})
        ]
        for integration in default_integrations:
            await db.integrations.insert_one(integration.model_dump())
        logger.info(f"Seeded {len(default_integrations)} integrations")
    
    # Start background task for scheduled alert checks
    asyncio.create_task(scheduled_alert_checks())
    logger.info("Started scheduled alert checks background task")

async def scheduled_alert_checks():
    """Run alert checks daily"""
    while True:
        try:
            logger.info("Running scheduled alert checks...")
            
            # Run the three alert check functions
            # We need to simulate an admin user context
            admin = await db.users.find_one({"role": "cs_lead"}, {"_id": 0})
            if admin:
                # Engagement gaps
                contacts = await db.client_contacts.find({}, {"_id": 0}).to_list(10000)
                now = datetime.now(timezone.utc)
                alerts_created = 0
                
                default_threshold = await db.alert_thresholds.find_one({"client_id": None}, {"_id": 0})
                if not default_threshold:
                    default_threshold = {"engagement_gap_decision_maker_days": 14, "engagement_gap_influencer_days": 30}
                
                for contact in contacts:
                    if contact.get('role_type') not in ['decision-maker', 'influencer']:
                        continue
                    
                    threshold = await db.alert_thresholds.find_one({"client_id": contact.get('client_id')}, {"_id": 0})
                    if not threshold:
                        threshold = default_threshold
                    
                    days_threshold = (
                        threshold.get('engagement_gap_decision_maker_days', 14)
                        if contact.get('role_type') == 'decision-maker'
                        else threshold.get('engagement_gap_influencer_days', 30)
                    )
                    
                    last_dates = []
                    for field in ['last_email_date', 'last_call_date', 'last_meeting_date']:
                        if contact.get(field):
                            try:
                                last_dates.append(datetime.fromisoformat(contact[field].replace('Z', '+00:00')))
                            except:
                                pass
                    
                    if not last_dates:
                        days_since = 999
                    else:
                        most_recent = max(last_dates)
                        days_since = (now - most_recent).days
                    
                    if days_since > days_threshold:
                        existing = await db.alerts.find_one({
                            "client_id": contact.get('client_id'),
                            "alert_type": "engagement_gap",
                            "status": {"$in": ["active", "acknowledged"]},
                            "trigger_data.contact_id": contact.get('id')
                        }, {"_id": 0})
                        
                        if not existing:
                            client = await db.clients.find_one({"id": contact.get('client_id')}, {"_id": 0})
                            client_name = client.get('name', 'Unknown') if client else 'Unknown'
                            
                            alert = Alert(
                                client_id=contact.get('client_id', ''),
                                client_name=client_name,
                                alert_type="engagement_gap",
                                severity="high" if contact.get('role_type') == 'decision-maker' else "medium",
                                title=f"No contact with {contact.get('name', 'Unknown')} ({contact.get('role_type', 'contact')}) in {days_since} days",
                                description=f"Contact {contact.get('name', 'Unknown')} has not been engaged in {days_since} days.",
                                trigger_data={
                                    "contact_id": contact.get('id'),
                                    "contact_name": contact.get('name'),
                                    "role_type": contact.get('role_type'),
                                    "days_since_contact": days_since
                                }
                            )
                            await db.alerts.insert_one(alert.model_dump())
                            alerts_created += 1
                
                logger.info(f"Created {alerts_created} engagement gap alerts")
                
                # Renewal alerts
                clients = await db.clients.find({"status": "active"}, {"_id": 0}).to_list(10000)
                renewal_alerts = 0
                
                for client in clients:
                    if not client.get('contract_end'):
                        continue
                    
                    try:
                        contract_end = datetime.fromisoformat(client['contract_end'].replace('Z', '+00:00'))
                        days_until = (contract_end - now).days
                    except:
                        continue
                    
                    threshold = await db.alert_thresholds.find_one({"client_id": client['id']}, {"_id": 0})
                    if not threshold:
                        threshold = default_threshold
                    
                    if 55 <= days_until <= 65 and threshold.get('renewal_alert_60_days', True):
                        existing = await db.alerts.find_one({
                            "client_id": client['id'],
                            "alert_type": "renewal",
                            "status": {"$in": ["active", "acknowledged"]},
                            "trigger_data.days_threshold": 60
                        }, {"_id": 0})
                        if not existing:
                            alert = Alert(
                                client_id=client['id'],
                                client_name=client['name'],
                                alert_type="renewal",
                                severity="medium",
                                title=f"Contract renewal in {days_until} days",
                                description=f"Contract expires on {client['contract_end']}",
                                trigger_data={"days_until_renewal": days_until, "days_threshold": 60}
                            )
                            await db.alerts.insert_one(alert.model_dump())
                            renewal_alerts += 1
                    
                    if 25 <= days_until <= 35 and threshold.get('renewal_alert_30_days', True):
                        existing = await db.alerts.find_one({
                            "client_id": client['id'],
                            "alert_type": "renewal",
                            "status": {"$in": ["active", "acknowledged"]},
                            "trigger_data.days_threshold": 30
                        }, {"_id": 0})
                        if not existing:
                            alert = Alert(
                                client_id=client['id'],
                                client_name=client['name'],
                                alert_type="renewal",
                                severity="high",
                                title=f"Contract renewal in {days_until} days",
                                description=f"Contract expires on {client['contract_end']}",
                                trigger_data={"days_until_renewal": days_until, "days_threshold": 30}
                            )
                            await db.alerts.insert_one(alert.model_dump())
                            renewal_alerts += 1
                    
                    if 10 <= days_until <= 20 and threshold.get('renewal_alert_15_days', True):
                        existing = await db.alerts.find_one({
                            "client_id": client['id'],
                            "alert_type": "renewal",
                            "status": {"$in": ["active", "acknowledged"]},
                            "trigger_data.days_threshold": 15
                        }, {"_id": 0})
                        if not existing:
                            alert = Alert(
                                client_id=client['id'],
                                client_name=client['name'],
                                alert_type="renewal",
                                severity="critical",
                                title=f"Contract renewal in {days_until} days - URGENT",
                                description=f"Contract expires on {client['contract_end']}",
                                trigger_data={"days_until_renewal": days_until, "days_threshold": 15}
                            )
                            await db.alerts.insert_one(alert.model_dump())
                            renewal_alerts += 1
                            
                            # Auto-create QBR scheduling follow-up when 60-day renewal fires
                            if days_until <= 65 and days_until >= 55:
                                existing_qbr = await db.followups.find_one({
                                    "client_id": client['id'],
                                    "title": {"$regex": "QBR", "$options": "i"},
                                    "status": "open"
                                }, {"_id": 0})
                                if not existing_qbr:
                                    qbr_followup = FollowUpItem(
                                        client_id=client['id'],
                                        client_name=client['name'],
                                        title=f"Schedule QBR for {client['name']} - Renewal in {days_until} days",
                                        description=f"Contract expires {client['contract_end']}. Schedule QBR 3-4 weeks before renewal to review performance and discuss terms.",
                                        priority_level=4,
                                        action_type="call_required",
                                        assigned_to=client.get('assigned_csm'),
                                        source="system_generated"
                                    )
                                    await db.followups.insert_one(qbr_followup.model_dump())
                                    logger.info(f"Auto-created QBR follow-up for {client['name']}")
                
                logger.info(f"Created {renewal_alerts} renewal alerts")
                
                # Innovation update due alerts
                innovation_alerts = 0
                for client in clients:
                    last_briefing = client.get('last_innovation_briefing')
                    if last_briefing:
                        try:
                            briefing_date = datetime.fromisoformat(last_briefing.replace('Z', '+00:00'))
                            if briefing_date.tzinfo is None:
                                briefing_date = briefing_date.replace(tzinfo=timezone.utc)
                            days_since = (now - briefing_date).days
                        except:
                            days_since = 999
                    else:
                        days_since = 999  # Never briefed
                    
                    if days_since > 60:
                        existing = await db.alerts.find_one({
                            "client_id": client['id'],
                            "alert_type": "innovation_update_due",
                            "status": {"$in": ["active", "acknowledged"]}
                        }, {"_id": 0})
                        if not existing:
                            alert = Alert(
                                client_id=client['id'],
                                client_name=client['name'],
                                alert_type="innovation_update_due",
                                severity="low",
                                title=f"Innovation update due for {client['name']} ({days_since} days since last briefing)",
                                description=f"Client has not been briefed on Anka product innovations in {days_since} days.",
                                trigger_data={"days_since_briefing": days_since, "last_briefing": last_briefing}
                            )
                            await db.alerts.insert_one(alert.model_dump())
                            innovation_alerts += 1
                logger.info(f"Created {innovation_alerts} innovation update due alerts")
            
            logger.info("Scheduled alert checks completed")
        except Exception as e:
            logger.error(f"Error in scheduled alert checks: {e}")
        
        # Wait 24 hours
        await asyncio.sleep(86400)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
